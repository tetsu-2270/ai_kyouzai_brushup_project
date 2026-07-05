from docx import Document

from src.docx_renderer import write_docx
from src.lesson_pages import build_lesson_document
from src.models import project_from_dict


def test_write_docx_contains_expected_content(tmp_path):
    document = build_lesson_document(project_from_dict({
        "project_title": "テスト教材",
        "target_reader": "テスター",
        "pages": [
            {
                "page_no": 1,
                "source_image": "a.png",
                "title": "P1",
                "summary": "概要文",
                "lines": [{"speaker": "まじょこ", "text": "こんにちは"}],
            }
        ],
    }))

    output_path = tmp_path / "nested" / "brushup.docx"
    write_docx(output_path, document)

    assert output_path.exists()

    saved_document = Document(str(output_path))
    texts = [p.text for p in saved_document.paragraphs]
    assert "テスト教材" in texts
    assert any("Page 1: P1" in t for t in texts)
    assert "概要文" in texts
    assert "まじょこ: こんにちは" in texts


def test_write_docx_does_not_expose_source_page_no_or_role(tmp_path):
    document = build_lesson_document(project_from_dict({
        "project_title": "テスト教材",
        "target_reader": "テスター",
        "pages": [
            {"page_no": 1, "source_image": "a.png", "title": "P1", "summary": "概要文", "lines": []},
        ],
    }))
    document.pages[0].source_page_no = [1, 2]
    document.pages[0].role = "explanation"

    output_path = tmp_path / "brushup.docx"
    write_docx(output_path, document)

    joined_text = "\n".join(p.text for p in Document(str(output_path)).paragraphs)
    assert "source_page_no" not in joined_text
    assert "explanation" not in joined_text
