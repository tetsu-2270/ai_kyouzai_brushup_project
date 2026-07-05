from __future__ import annotations

from pathlib import Path

from docx import Document

from .lesson_pages import LessonDocument, parse_body_lines


def render_docx(document: LessonDocument) -> Document:
    doc = Document()

    doc.add_heading(document.project_title, level=0)
    doc.add_paragraph(f"対象読者: {document.target_reader}")

    doc.add_heading("全体方針", level=1)
    for point in [
        "ページ単位で情報を整理する。",
        "話者ごとに台詞を分ける。",
        "教材として理解しやすい順序・表現に整える。",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    for page in document.pages:
        doc.add_heading(f"Page {page.page_no}: {page.title}", level=1)

        doc.add_heading("概要", level=2)
        doc.add_paragraph(page.summary or "未設定")

        doc.add_heading("本文", level=2)
        parsed = parse_body_lines(page.body)
        if parsed:
            for speaker, text in parsed:
                label = f"{speaker}: {text}" if speaker else text
                doc.add_paragraph(label, style="List Bullet")
        else:
            doc.add_paragraph("未設定")

    return doc


def write_docx(path: str | Path, document: LessonDocument) -> None:
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    render_docx(document).save(output_path)
